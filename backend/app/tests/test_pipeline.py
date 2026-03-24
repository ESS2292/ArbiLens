from io import BytesIO

from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from reportlab.pdfgen import canvas
from sqlalchemy.orm import Session, sessionmaker

from app.core.config import get_settings
from app.models.analysis_job import AnalysisJob
from app.models.clause import Clause
from app.models.document import Document
from app.models.document_version import DocumentVersion
from app.models.enums import ClauseType, DocumentStatus, JobStatus, OrganizationRole, RiskScope, RiskSeverity
from app.models.organization import Organization
from app.models.risk import Risk
from app.models.user import User
from app.tasks import document_tasks
from app.tests.conftest import FakeObjectStorageService, register_user


def _pdf_bytes() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(72, 800, "Limitation of Liability")
    pdf.drawString(72, 782, "Provider liability is unlimited and without limitation.")
    pdf.drawString(72, 764, "Termination")
    pdf.drawString(72, 746, "Provider may terminate for convenience at any time.")
    pdf.save()
    return buffer.getvalue()


def _docx_bytes() -> bytes:
    buffer = BytesIO()
    document = DocxDocument()
    document.add_heading("Limitation of Liability", level=1)
    document.add_paragraph("Provider liability is unlimited and without limitation.")
    document.add_heading("Confidentiality", level=1)
    document.add_paragraph("Each party shall protect confidential information.")
    document.save(buffer)
    return buffer.getvalue()


def test_pipeline_happy_path_runs_from_uploaded_document_to_completed_results(
    client: TestClient,
    db_session: Session,
    monkeypatch,
) -> None:
    monkeypatch.setenv("MAX_UPLOAD_SIZE_BYTES", "65536")
    get_settings.cache_clear()
    registered = register_user(client)
    token = registered["access_token"]

    upload_response = client.post(
        "/api/v1/documents/upload",
        headers={"Authorization": f"Bearer {token}"},
        files={"upload": ("contract.pdf", _pdf_bytes(), "application/pdf")},
    )
    assert upload_response.status_code == 201
    payload = upload_response.json()

    version = db_session.get(DocumentVersion, payload["document_version_id"])
    assert version is not None
    version.original_filename = "msa.pdf"
    db_session.commit()

    task_session_factory = sessionmaker(bind=db_session.get_bind(), autocommit=False, autoflush=False, class_=Session)

    monkeypatch.setattr(document_tasks, "get_session_factory", lambda: task_session_factory)
    monkeypatch.setattr(document_tasks, "get_object_storage_service", lambda: client.storage)

    document_tasks.parse_document_task.run(payload["job_id"])
    document_tasks.normalize_document_task.run(payload["job_id"])
    document_tasks.extract_clauses_task.run(payload["job_id"])
    document_tasks.analyze_risks_task.run(payload["job_id"])

    db_session.expire_all()
    refreshed_job = db_session.get(AnalysisJob, payload["job_id"])
    refreshed_version = db_session.get(DocumentVersion, payload["document_version_id"])
    refreshed_document = db_session.get(Document, payload["document_id"])

    assert refreshed_job is not None
    assert refreshed_job.status == JobStatus.completed
    assert refreshed_job.current_stage == "completed"
    assert refreshed_version is not None
    assert refreshed_version.extracted_text is not None
    assert len(refreshed_version.chunks) >= 1
    assert len(refreshed_version.clauses) >= 1
    assert len(refreshed_version.risks) >= 1
    assert all(risk.citations for risk in refreshed_version.risks)
    assert refreshed_document is not None
    assert refreshed_document.status == DocumentStatus.completed

    status_response = client.get(
        f"/api/v1/documents/{payload['document_id']}/status",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert status_response.status_code == 200
    assert status_response.json()["document_status"] == "completed"

    summary_response = client.get(
        f"/api/v1/documents/{payload['document_id']}/summary",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert summary_response.status_code == 200
    assert summary_response.json()["overall_risk_score"] >= 0


def test_pipeline_happy_path_runs_for_docx_upload(
    client: TestClient,
    db_session: Session,
    monkeypatch,
) -> None:
    monkeypatch.setenv("MAX_UPLOAD_SIZE_BYTES", "65536")
    get_settings.cache_clear()
    registered = register_user(client)
    token = registered["access_token"]

    upload_response = client.post(
        "/api/v1/documents/upload",
        headers={"Authorization": f"Bearer {token}"},
        files={
            "upload": (
                "contract.docx",
                _docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 201
    payload = upload_response.json()

    task_session_factory = sessionmaker(bind=db_session.get_bind(), autocommit=False, autoflush=False, class_=Session)

    monkeypatch.setattr(document_tasks, "get_session_factory", lambda: task_session_factory)
    monkeypatch.setattr(document_tasks, "get_object_storage_service", lambda: client.storage)

    document_tasks.parse_document_task.run(payload["job_id"])
    document_tasks.normalize_document_task.run(payload["job_id"])
    document_tasks.extract_clauses_task.run(payload["job_id"])
    document_tasks.analyze_risks_task.run(payload["job_id"])

    db_session.expire_all()
    refreshed_version = db_session.get(DocumentVersion, payload["document_version_id"])
    refreshed_document = db_session.get(Document, payload["document_id"])

    assert refreshed_version is not None
    assert refreshed_version.extracted_text is not None
    assert len(refreshed_version.chunks) >= 1
    assert len(refreshed_version.clauses) >= 1
    assert len(refreshed_version.risks) >= 1
    assert refreshed_document is not None
    assert refreshed_document.status == DocumentStatus.completed


def test_pipeline_parse_failure_marks_job_and_document_failed(
    db_session: Session,
    monkeypatch,
) -> None:
    organization = Organization(name="Acme", slug="acme-pipeline-fail")
    user = User(
        organization=organization,
        email="owner@example.com",
        full_name="Owner",
        password_hash="hashed",
        role=OrganizationRole.owner,
        is_active=True,
    )
    document = Document(
        organization=organization,
        created_by_user=user,
        filename="contract.pdf",
        status=DocumentStatus.queued,
        latest_version_number=1,
    )
    version = DocumentVersion(
        document=document,
        uploaded_by_user=user,
        version_number=1,
        storage_key="documents/bad/contract.pdf",
        original_filename="contract.pdf",
        content_type="application/pdf",
        file_extension=".pdf",
        file_size_bytes=17,
        sha256_hash="a" * 64,
        is_current=True,
        status=DocumentStatus.queued,
    )
    job = AnalysisJob(
        organization=organization,
        document=document,
        document_version=version,
        requested_by_user=user,
        status=JobStatus.queued,
        current_stage="queued",
    )
    db_session.add_all([organization, user, document, version, job])
    db_session.commit()

    storage = FakeObjectStorageService()
    storage.uploaded_objects[version.storage_key] = b"not a real pdf document"
    task_session_factory = sessionmaker(bind=db_session.get_bind(), autocommit=False, autoflush=False, class_=Session)

    monkeypatch.setattr(document_tasks, "get_session_factory", lambda: task_session_factory)
    monkeypatch.setattr(document_tasks, "get_object_storage_service", lambda: storage)

    document_tasks.parse_document_task.run(str(job.id))

    db_session.expire_all()
    refreshed_job = db_session.get(AnalysisJob, job.id)
    refreshed_document = db_session.get(Document, document.id)

    assert refreshed_job is not None
    assert refreshed_job.status == JobStatus.failed
    assert refreshed_job.error_stage == "parsing"
    assert refreshed_job.error_code in {"pdf_unreadable", "pipeline_error"}
    assert refreshed_document is not None
    assert refreshed_document.status == DocumentStatus.failed


def test_pipeline_re_running_completed_analysis_is_idempotent(
    db_session: Session,
    monkeypatch,
) -> None:
    organization = Organization(name="Acme", slug="acme-idempotent")
    user = User(
        organization=organization,
        email="idempotent@example.com",
        full_name="Owner",
        password_hash="hashed",
        role=OrganizationRole.owner,
        is_active=True,
    )
    document = Document(
        organization=organization,
        created_by_user=user,
        filename="contract.pdf",
        status=DocumentStatus.completed,
        latest_version_number=1,
    )
    version = DocumentVersion(
        document=document,
        uploaded_by_user=user,
        version_number=1,
        storage_key="documents/idempotent/contract.pdf",
        original_filename="contract.pdf",
        content_type="application/pdf",
        file_extension=".pdf",
        file_size_bytes=128,
        sha256_hash="b" * 64,
        is_current=True,
        status=DocumentStatus.completed,
    )
    job = AnalysisJob(
        organization=organization,
        document=document,
        document_version=version,
        requested_by_user=user,
        status=JobStatus.completed,
        current_stage="completed",
    )
    clause = Clause(
        document_version=version,
        clause_type=ClauseType.confidentiality,
        title="Confidentiality",
        text="Confidential information must be protected.",
        normalized_text="Confidential information must be protected.",
        confidence=0.9,
        source_method="heuristic",
        page_start=1,
        page_end=1,
    )
    risk = Risk(
        document=document,
        document_version=version,
        clause=clause,
        analysis_job=job,
        scope=RiskScope.clause,
        severity=RiskSeverity.medium,
        category="confidentiality",
        title="Existing risk",
        summary="Existing risk",
        score=30,
        rationale="Already analyzed.",
        recommendation="No-op",
        confidence=0.8,
        citations=[{"reference_type": "clause", "clause_id": "existing"}],
        deterministic_rule_code="existing",
        evidence_text=clause.text,
    )
    db_session.add_all([organization, user, document, version, job, clause, risk])
    db_session.commit()

    task_session_factory = sessionmaker(bind=db_session.get_bind(), autocommit=False, autoflush=False, class_=Session)
    monkeypatch.setattr(document_tasks, "get_session_factory", lambda: task_session_factory)

    document_tasks.analyze_risks_task.run(str(job.id))

    db_session.expire_all()
    refreshed_version = db_session.get(DocumentVersion, version.id)
    assert refreshed_version is not None
    assert len(refreshed_version.risks) == 1
