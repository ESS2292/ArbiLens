from io import BytesIO

from docx import Document as DocxDocument
from reportlab.pdfgen import canvas

from app.services.parsing.docx_parser import DocxParsingService
from app.services.parsing.pdf_parser import PdfParsingService


def _build_pdf_bytes(lines: list[str]) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    y = 800
    for line in lines:
        pdf.drawString(72, y, line)
        y -= 18
    pdf.save()
    return buffer.getvalue()


def _build_docx_bytes() -> bytes:
    buffer = BytesIO()
    document = DocxDocument()
    document.add_heading("1. Introduction", level=1)
    document.add_paragraph("This agreement is entered into by and between the parties.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Term"
    table.rows[0].cells[1].text = "12 months"
    table.rows[1].cells[0].text = "Renewal"
    table.rows[1].cells[1].text = "Automatic"
    document.save(buffer)
    return buffer.getvalue()


def test_pdf_parser_extracts_text_and_metadata() -> None:
    content = _build_pdf_bytes(["MASTER SERVICES AGREEMENT", "Section 1. Scope", "Payment due in 30 days."])

    result = PdfParsingService().parse(content)

    assert result.parser_used == "pypdf"
    assert result.page_count == 1
    assert result.ocr_needed is False
    assert "MASTER SERVICES AGREEMENT" in result.full_text
    assert result.metadata["format"] == "pdf"


def test_docx_parser_preserves_headings_and_tables() -> None:
    result = DocxParsingService().parse(_build_docx_bytes())

    assert result.parser_used == "python-docx"
    assert "This agreement is entered into" in result.full_text
    assert any(section.heading == "1. Introduction" for section in result.sections)
    assert any("Renewal | Automatic" in block["text"] for block in result.metadata["blocks"])
